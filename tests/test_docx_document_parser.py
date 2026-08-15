from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

import pytest
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table

from filelore.documents import DocxDocumentParser, TextBlockType
from filelore.metadata import DocumentFormat


def create_text_docx(path: Path) -> None:
    document = Document()
    _configure_fixture_styles(document)
    core = document.core_properties
    core.title = "Türkçe DOCX Belgesi"
    core.author = "Örnek Yazar"
    core.subject = "Yapısal metin çıkarma"
    core.keywords = "Türkçe, arama, İstanbul"
    core.language = "tr"
    core.created = datetime(2024, 1, 2, 10, 30)
    core.modified = datetime(2024, 1, 3, 11, 45)

    document.add_heading("Ana Başlık", level=1)
    document.add_paragraph("Türkçe, 日本語 ve العربية aynı paragrafta kalır.")
    document.add_heading("Ayrıntılar", level=2)
    document.add_paragraph("Birinci öğe", style="List Bullet")
    document.add_paragraph("İkinci öğe", style="List Number")
    document.add_paragraph("Alıntı metni", style="Quote")

    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(1.875)
    table.columns[1].width = Inches(4.625)
    table.cell(0, 0).text = "Alan"
    table.cell(0, 1).text = "Değer"
    table.cell(1, 0).text = "Dil"
    table.cell(1, 1).text = "Türkçe"
    _set_table_width_and_indent(table, width_dxa=9360, indent_dxa=120)

    document.sections[0].header.paragraphs[0].text = "Dahil edilmeyen üst bilgi"
    document.sections[0].footer.paragraphs[0].text = "Dahil edilmeyen alt bilgi"
    document.save(path)


def _configure_fixture_styles(document: DocxDocument) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        "Heading 1": (16, "2E74B5", 16, 8),
        "Heading 2": (13, "2E74B5", 12, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.167


def _set_table_width_and_indent(
    table: Table,
    *,
    width_dxa: int,
    indent_dxa: int,
) -> None:
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(width_dxa))

    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), str(indent_dxa))


def test_docx_parser_extracts_structure_metadata_and_unicode(
    tmp_path: Path,
) -> None:
    path = tmp_path / "belge.DOCX"
    create_text_docx(path)

    document = DocxDocumentParser().parse(path)

    assert document.metadata.document_format is DocumentFormat.DOCX
    assert document.metadata.extension == ".docx"
    assert document.metadata.mime_type.endswith("wordprocessingml.document")
    assert document.metadata.title == "Türkçe DOCX Belgesi"
    assert document.metadata.authors == ("Örnek Yazar",)
    assert document.metadata.language == "tr"
    assert document.metadata.created_at == datetime(
        2024, 1, 2, 10, 30, tzinfo=timezone.utc
    )
    assert document.metadata.content_modified_at == datetime(
        2024, 1, 3, 11, 45, tzinfo=timezone.utc
    )
    assert document.metadata.page_count is None
    assert document.metadata.properties["subject"] == "Yapısal metin çıkarma"
    assert document.metadata.properties["keywords"] == (
        "Türkçe",
        "arama",
        "İstanbul",
    )
    assert document.metadata.properties["section_count"] == 1
    assert [block.block_type for block in document.blocks] == [
        TextBlockType.HEADING,
        TextBlockType.PARAGRAPH,
        TextBlockType.HEADING,
        TextBlockType.LIST_ITEM,
        TextBlockType.LIST_ITEM,
        TextBlockType.QUOTE,
        TextBlockType.TABLE,
    ]
    assert document.blocks[1].text == (
        "Türkçe, 日本語 ve العربية aynı paragrafta kalır."
    )
    assert document.blocks[1].location.section_path == ("Ana Başlık",)
    assert document.blocks[3].location.section_path == (
        "Ana Başlık",
        "Ayrıntılar",
    )
    assert document.blocks[3].attributes["list_level"] == 0
    assert document.blocks[5].text == "Alıntı metni"
    assert document.blocks[6].text == "Alan | Değer\nDil | Türkçe"
    assert document.blocks[6].attributes["rows"] == 2
    assert all("üst bilgi" not in block.text for block in document.blocks)
    assert all("alt bilgi" not in block.text for block in document.blocks)


def test_docx_parser_uses_a_styled_title_fallback(tmp_path: Path) -> None:
    path = tmp_path / "title.docx"
    source = Document()
    source.add_paragraph("Stil Başlığı", style="Title")
    source.add_paragraph("İçerik")
    source.save(path)

    document = DocxDocumentParser().parse(path)

    assert document.metadata.title == "Stil Başlığı"


def test_docx_parser_allows_an_empty_body(tmp_path: Path) -> None:
    path = tmp_path / "empty.docx"
    Document().save(path)

    document = DocxDocumentParser().parse(path)

    assert document.blocks == ()


def test_docx_parser_enforces_archive_limits(tmp_path: Path) -> None:
    path = tmp_path / "limited.docx"
    create_text_docx(path)

    with pytest.raises(ValueError, match="uncompressed size limit"):
        DocxDocumentParser(max_uncompressed_bytes=1).parse(path)
    with pytest.raises(ValueError, match="more than 1 members"):
        DocxDocumentParser(max_archive_members=1).parse(path)


def test_docx_parser_rejects_corrupt_unsupported_and_missing_files(
    tmp_path: Path,
) -> None:
    corrupt_path = tmp_path / "corrupt.docx"
    corrupt_path.write_text("not a DOCX", encoding="utf-8")
    parser = DocxDocumentParser()

    with pytest.raises(ValueError, match="Could not parse DOCX"):
        parser.parse(corrupt_path)
    with pytest.raises(ValueError, match="Unsupported document extension"):
        parser.parse(tmp_path / "notes.txt")
    with pytest.raises(FileNotFoundError):
        parser.parse(tmp_path / "missing.docx")


@pytest.mark.parametrize(
    ("uncompressed_limit", "member_limit"),
    ((0, 1), (1, 0)),
)
def test_docx_parser_rejects_invalid_archive_limits(
    uncompressed_limit: int,
    member_limit: int,
) -> None:
    with pytest.raises(ValueError, match="limit must be positive"):
        DocxDocumentParser(
            max_uncompressed_bytes=uncompressed_limit,
            max_archive_members=member_limit,
        )
