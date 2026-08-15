"""Structure-aware DOCX parser backed by python-docx."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.document import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml.etree import XMLSyntaxError

from filelore.documents.models import (
    ParsedDocument,
    SourceLocation,
    TextBlock,
    TextBlockType,
)
from filelore.documents.parsers.base import DocumentParser
from filelore.metadata import DocumentFormat, DocumentMetadata


_HEADING_STYLE = re.compile(r"^heading\s*([1-9])$", re.IGNORECASE)


class DocxDocumentParser(DocumentParser):
    """Extract body structure and core properties from Word documents."""

    supported_extensions = frozenset({".docx"})
    default_max_uncompressed_bytes = 256 * 1024 * 1024
    default_max_archive_members = 10_000

    def __init__(
        self,
        *,
        max_uncompressed_bytes: int = default_max_uncompressed_bytes,
        max_archive_members: int = default_max_archive_members,
    ) -> None:
        if max_uncompressed_bytes <= 0:
            raise ValueError("DOCX uncompressed size limit must be positive")
        if max_archive_members <= 0:
            raise ValueError("DOCX archive member limit must be positive")
        self.max_uncompressed_bytes = max_uncompressed_bytes
        self.max_archive_members = max_archive_members

    def parse(self, path: str | Path) -> ParsedDocument:
        document_path = self.prepare_path(path)
        try:
            _validate_archive(
                document_path,
                max_uncompressed_bytes=self.max_uncompressed_bytes,
                max_archive_members=self.max_archive_members,
            )
        except BadZipFile as error:
            raise ValueError(
                f"Could not parse DOCX document: {document_path}"
            ) from error

        try:
            document = Document(document_path)
            blocks = _extract_blocks(document)
            metadata = _document_metadata(document, document_path, blocks)
            return ParsedDocument(metadata=metadata, blocks=blocks)
        except (
            BadZipFile,
            PackageNotFoundError,
            XMLSyntaxError,
            KeyError,
            TypeError,
            ValueError,
        ) as error:
            raise ValueError(
                f"Could not parse DOCX document: {document_path}"
            ) from error


def _validate_archive(
    path: Path,
    *,
    max_uncompressed_bytes: int,
    max_archive_members: int,
) -> None:
    with ZipFile(path) as archive:
        members = archive.infolist()
        if len(members) > max_archive_members:
            raise ValueError(
                f"DOCX archive contains more than {max_archive_members} members"
            )
        uncompressed_size = sum(member.file_size for member in members)
        if uncompressed_size > max_uncompressed_bytes:
            raise ValueError(
                "DOCX archive exceeds the "
                f"{max_uncompressed_bytes}-byte uncompressed size limit"
            )
        if any(member.flag_bits & 0x1 for member in members):
            raise ValueError("Encrypted DOCX archives are not supported")

        names = {member.filename for member in members}
        required = {"[Content_Types].xml", "word/document.xml"}
        if not required.issubset(names):
            raise ValueError("DOCX archive is missing required document parts")


def _extract_blocks(document: DocxDocument) -> tuple[TextBlock, ...]:
    blocks: list[TextBlock] = []
    section_stack: list[tuple[int, str]] = []

    for item in document.iter_inner_content():
        if isinstance(item, Paragraph):
            text = _normalize_text(item.text)
            if not text:
                continue
            style_id = _style_id(item)
            style_name = _style_name(item)
            heading_level = _heading_level(style_id, style_name)
            attributes: dict[str, Any] = {}
            if style_name:
                attributes["style"] = style_name

            if heading_level is not None:
                while section_stack and section_stack[-1][0] >= heading_level:
                    section_stack.pop()
                section_stack.append((heading_level, text))
                block_type = TextBlockType.HEADING
                attributes["level"] = heading_level
            else:
                list_attributes = _list_attributes(item, style_id)
                if list_attributes is not None:
                    block_type = TextBlockType.LIST_ITEM
                    attributes.update(list_attributes)
                elif style_id.casefold() in {"quote", "intensequote"}:
                    block_type = TextBlockType.QUOTE
                else:
                    block_type = TextBlockType.PARAGRAPH

            blocks.append(
                TextBlock(
                    index=len(blocks),
                    block_type=block_type,
                    text=text,
                    location=_location(section_stack),
                    attributes=attributes,
                )
            )
        elif isinstance(item, Table):
            text = _table_text(item)
            if not text:
                continue
            attributes = {
                "rows": len(item.rows),
                "columns": len(item.columns),
            }
            if item.style is not None and item.style.name:
                attributes["style"] = item.style.name
            blocks.append(
                TextBlock(
                    index=len(blocks),
                    block_type=TextBlockType.TABLE,
                    text=text,
                    location=_location(section_stack),
                    attributes=attributes,
                )
            )

    return tuple(blocks)


def _document_metadata(
    document: DocxDocument,
    path: Path,
    blocks: tuple[TextBlock, ...],
) -> DocumentMetadata:
    core = document.core_properties
    title = _first_text(
        core.title,
        _styled_title(document),
        _first_level_one_heading(blocks),
    )
    author = _first_text(core.author)
    stat = path.stat()

    return DocumentMetadata(
        path=path,
        extension=path.suffix.casefold(),
        mime_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        size_bytes=stat.st_size,
        modified_at=datetime.fromtimestamp(stat.st_mtime).astimezone(),
        document_format=DocumentFormat.DOCX,
        title=title,
        authors=(author,) if author else (),
        created_at=core.created if isinstance(core.created, datetime) else None,
        content_modified_at=(
            core.modified if isinstance(core.modified, datetime) else None
        ),
        language=_first_text(core.language),
        properties=_document_properties(document),
    )


def _document_properties(document: DocxDocument) -> dict[str, Any]:
    core = document.core_properties
    properties: dict[str, Any] = {"section_count": len(document.sections)}
    for key in (
        "category",
        "content_status",
        "identifier",
        "subject",
        "version",
    ):
        value = _first_text(getattr(core, key, None))
        if value:
            properties[key] = value

    keywords = _keywords(core.keywords)
    if keywords:
        properties["keywords"] = keywords
    if isinstance(core.revision, int):
        properties["revision"] = core.revision
    return properties


def _styled_title(document: DocxDocument) -> str | None:
    for paragraph in document.paragraphs:
        if _style_id(paragraph).casefold() == "title":
            title = _normalize_text(paragraph.text)
            if title:
                return title
    return None


def _first_level_one_heading(blocks: tuple[TextBlock, ...]) -> str | None:
    for block in blocks:
        if (
            block.block_type is TextBlockType.HEADING
            and block.attributes.get("level") == 1
        ):
            return block.text
    return None


def _style_id(paragraph: Paragraph) -> str:
    return paragraph.style.style_id if paragraph.style is not None else ""


def _style_name(paragraph: Paragraph) -> str:
    return paragraph.style.name if paragraph.style is not None else ""


def _heading_level(style_id: str, style_name: str) -> int | None:
    for value in (style_id, style_name):
        match = _HEADING_STYLE.fullmatch(value.replace("_", " "))
        if match:
            return int(match.group(1))
    return None


def _list_attributes(
    paragraph: Paragraph,
    style_id: str,
) -> dict[str, int] | None:
    number_properties = None
    paragraph_properties = paragraph._p.pPr
    if paragraph_properties is not None:
        number_properties = paragraph_properties.numPr
    if number_properties is None and paragraph.style is not None:
        style_properties = paragraph.style._element.pPr
        if style_properties is not None:
            number_properties = style_properties.numPr

    style_is_list = style_id.casefold().startswith(("listbullet", "listnumber"))
    if number_properties is None and not style_is_list:
        return None

    attributes: dict[str, int] = {"list_level": 0}
    if number_properties is not None:
        if number_properties.ilvl is not None:
            attributes["list_level"] = int(number_properties.ilvl.val)
        if number_properties.numId is not None:
            attributes["numbering_id"] = int(number_properties.numId.val)
    return attributes


def _table_text(table: Table) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = [_single_line_text(cell.text) for cell in row.cells]
        if any(cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def _location(section_stack: list[tuple[int, str]]) -> SourceLocation:
    return SourceLocation(
        section_path=tuple(title for _, title in section_stack)
    )


def _normalize_text(value: str) -> str:
    normalized = value.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    lines = (
        _single_line_text(line)
        for line in normalized.splitlines()
    )
    return "\n".join(line for line in lines if line)


def _single_line_text(value: str) -> str:
    return " ".join(value.replace("\x00", "").split())


def _first_text(*values: Any) -> str | None:
    for value in values:
        if value is None:
            continue
        text = _single_line_text(str(value))
        if text:
            return text
    return None


def _keywords(value: Any) -> tuple[str, ...]:
    if value is None:
        return ()
    prepared: list[str] = []
    seen: set[str] = set()
    for candidate in re.split(r"[,;]", str(value)):
        keyword = _first_text(candidate)
        if keyword is None or keyword.casefold() in seen:
            continue
        seen.add(keyword.casefold())
        prepared.append(keyword)
    return tuple(prepared)
